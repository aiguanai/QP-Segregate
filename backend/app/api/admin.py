from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form, BackgroundTasks, Query
from sqlalchemy.orm import Session
from sqlalchemy import func
from typing import List, Optional
from datetime import datetime
import os
import shutil
import sys
from app.core.database import get_db
from app.models.user import User
from app.models.question_paper import QuestionPaper, ProcessingStatus, ExamType, SemesterType
from app.models.question import Question, ReviewQueue, ReviewStatus, BloomLevel, BloomCategory
from app.models.course import Course
from app.api.auth import get_current_user
from app.services.ocr_service import OCRService
from app.services.classification_service import ClassificationService
from app.tasks.processing import process_question_paper
from pydantic import BaseModel, field_validator

router = APIRouter()

class FileUploadResponse(BaseModel):
    upload_id: str
    page_count: int
    file_size: int
    temp_path: str
    file_type: str  # 'pdf' or 'docx'
    original_filename: str

class MetadataSubmitRequest(BaseModel):
    upload_id: str
    course_code: str
    academic_year: int
    semester_type: str
    exam_type: Optional[str] = None  # Optional for syllabus
    exam_date: Optional[datetime] = None
    file_type: Optional[str] = "question_paper"  # 'question_paper' or 'syllabus'
    
    @field_validator('exam_date', mode='before')
    @classmethod
    def parse_exam_date(cls, v):
        if v is None or v == '':
            return None
        if isinstance(v, str):
            # Handle date string format (YYYY-MM-DD)
            if len(v) == 10 and v.count('-') == 2:
                try:
                    # Parse date and convert to datetime at midnight
                    from datetime import date as date_type
                    date_obj = date_type.fromisoformat(v)
                    return datetime.combine(date_obj, datetime.min.time())
                except ValueError:
                    # If parsing fails, try datetime format
                    try:
                        return datetime.fromisoformat(v.replace('Z', '+00:00'))
                    except ValueError:
                        pass
        return v

class ProcessingStatusResponse(BaseModel):
    paper_id: int
    status: str
    progress: float
    questions_extracted: int
    errors: List[str] = []

class ReviewQueueItem(BaseModel):
    review_id: int
    question_id: int
    question_text: str
    issue_type: str
    suggested_correction: dict
    priority: int
    created_at: datetime
    marks: Optional[int] = None
    bloom_level: Optional[int] = None
    bloom_category: Optional[str] = None
    unit_id: Optional[int] = None
    unit_name: Optional[str] = None
    unit_topics: Optional[List[str]] = None  # Topics from the unit
    course_code: Optional[str] = None
    course_name: Optional[str] = None
    topic_tags: Optional[List[str]] = None  # Topic tags assigned to the question
    image_path: Optional[str] = None
    
    class Config:
        from_attributes = True

class QuestionApprovalRequest(BaseModel):
    unit_id: Optional[int] = None
    bloom_level: Optional[int] = None
    marks: Optional[int] = None
    approved: bool = True

class QuestionUpdateRequest(BaseModel):
    unit_id: Optional[int] = None
    topic_tags: Optional[List[str]] = None
    marks: Optional[int] = None
    bloom_taxonomy_level: Optional[int] = None
    bloom_category: Optional[str] = None
    question_text: Optional[str] = None

def convert_docx_to_pdf(docx_path: str) -> str:
    """Convert DOCX file to PDF (optional - for storage purposes)"""
    try:
        from docx import Document
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        
        doc = Document(docx_path)
        pdf_path = docx_path.replace('.docx', '.pdf').replace('.doc', '.pdf')
        
        c = canvas.Canvas(pdf_path, pagesize=letter)
        y = 750
        for para in doc.paragraphs:
            if para.text.strip():
                # Handle long lines by wrapping
                text = para.text
                words = text.split()
                line = ""
                for word in words:
                    if len(line + word) < 80:
                        line += word + " "
                    else:
                        if line:
                            c.drawString(50, y, line.strip())
                            y -= 20
                        line = word + " "
                        if y < 50:
                            c.showPage()
                            y = 750
                if line:
                    c.drawString(50, y, line.strip())
                    y -= 20
                if y < 50:
                    c.showPage()
                    y = 750
        
        c.save()
        return pdf_path
    except Exception as e:
        # If conversion fails, we can still process DOCX directly
        # Just return the original path and let OCR service handle it
        return docx_path

def get_page_count(file_path: str, file_type: str) -> int:
    """Get page count for PDF or DOCX file. Returns 1 if unable to determine."""
    if file_type == 'pdf':
        # Try multiple methods for better accuracy
        # Method 1: Try pdfplumber (most reliable)
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                return len(pdf.pages)
        except ImportError:
            # pdfplumber not installed, try PyPDF2
            pass
        except Exception as e:
            # pdfplumber failed, try PyPDF2 as fallback
            import sys
            sys.stderr.write(f"⚠️  pdfplumber error: {str(e)}, trying PyPDF2...\n")
        
        # Method 2: Try PyPDF2
        try:
            import PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                return len(pdf_reader.pages)
        except ImportError:
            # PyPDF2 not installed, try pdf2image
            import sys
            sys.stderr.write("⚠️  PyPDF2 not available, trying pdf2image...\n")
        except Exception as e:
            # PyPDF2 failed, try pdf2image as fallback
            import sys
            sys.stderr.write(f"⚠️  PyPDF2 error: {str(e)}, trying pdf2image...\n")
        
        # Method 3: Try pdf2image (requires poppler)
        try:
            from pdf2image import convert_from_path
            images = convert_from_path(file_path)
            return len(images)
        except ImportError:
            # pdf2image not installed - use fallback
            import sys
            sys.stderr.write("⚠️  pdf2image not available, using fallback page count\n")
            return 1
        except Exception as e:
            # Poppler not installed or other error - use fallback
            import sys
            error_msg = str(e)
            if "poppler" in error_msg.lower() or "path" in error_msg.lower():
                sys.stderr.write(f"⚠️  Poppler not installed: {error_msg}. Using fallback page count.\n")
            else:
                sys.stderr.write(f"⚠️  Error reading PDF: {error_msg}. Using fallback page count.\n")
            sys.stderr.flush()
            return 1  # Return 1 as fallback instead of raising error
    elif file_type == 'docx':
        try:
            from docx import Document
            doc = Document(file_path)
            
            # Better page estimation for DOCX
            # Count paragraphs with actual content
            para_count = len([p for p in doc.paragraphs if p.text.strip()])
            
            # Count tables (each table can take significant space)
            table_count = len(doc.tables)
            
            # Count images/embedded objects
            image_count = 0
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_count += 1
            
            # Better estimation:
            # - Base: paragraphs (roughly 25-30 lines per page, ~3-4 paragraphs per line)
            # - Tables: each table roughly 0.5-1 page depending on size
            # - Images: each image roughly 0.3-0.5 pages
            estimated_pages = max(1, 
                (para_count // 8) +  # ~8 paragraphs per page
                (table_count * 1) +  # each table ~1 page
                (image_count * 1)    # each image ~1 page
            )
            
            return estimated_pages
        except ImportError:
            # python-docx not installed - use fallback
            import sys
            sys.stderr.write("⚠️  python-docx not available, using fallback page count\n")
            return 1
        except Exception as e:
            # Error reading DOCX - use fallback
            import sys
            sys.stderr.write(f"⚠️  Error reading DOCX: {str(e)}. Using fallback page count.\n")
            sys.stderr.flush()
            return 1  # Return 1 as fallback instead of raising error
    else:
        return 1

@router.post("/upload-file", response_model=FileUploadResponse)
async def upload_file(
    file: UploadFile = File(...),
    file_type: str = Form(...),  # 'question_paper' or 'syllabus'
    current_user: User = Depends(get_current_user)
):
    """Step 1: Upload PDF or DOCX file and return upload details"""
    import sys
    sys.stderr.write(f"📤 Upload request received\n")
    sys.stderr.write(f"   File: {file.filename}\n")
    sys.stderr.write(f"   File type (document): {file_type}\n")
    sys.stderr.write(f"   File type (received): {type(file_type)}, value: '{file_type}'\n")
    sys.stderr.flush()
    
    if current_user.role != "ADMIN":
        raise HTTPException(status_code=403, detail="Admin access required")
    
    # Validate file_type parameter
    if file_type not in ['question_paper', 'syllabus']:
        error_msg = f"Invalid file_type: '{file_type}'. Must be 'question_paper' or 'syllabus'"
        sys.stderr.write(f"❌ {error_msg}\n")
        sys.stderr.flush()
        raise HTTPException(
            status_code=400, 
            detail=error_msg
        )
    
    filename_lower = file.filename.lower() if file.filename else ""
    if not (filename_lower.endswith('.pdf') or filename_lower.endswith('.docx') or filename_lower.endswith('.doc')):
        raise HTTPException(status_code=400, detail="Only PDF and DOCX files are allowed")
    
    # Determine file type
    is_docx = filename_lower.endswith('.docx') or filename_lower.endswith('.doc')
    detected_file_type = 'docx' if is_docx else 'pdf'
    
    # Create upload ID and temp directory
    upload_id = f"upload_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{file.filename}"
    temp_dir = "tmp/uploads"
    os.makedirs(temp_dir, exist_ok=True)
    temp_path = os.path.join(temp_dir, upload_id)
    
    # Save file to temp location
    with open(temp_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    # For DOCX files, we'll keep the original and process it directly
    # No need to convert to PDF for processing (OCR service can handle DOCX)
    
    # Get file size
    file_size = os.path.getsize(temp_path)
    
    # Get page count
    try:
        page_count = get_page_count(temp_path, detected_file_type)
    except HTTPException:
        # Re-raise HTTPExceptions from get_page_count (they already have proper error messages)
        raise
    except Exception as e:
        error_msg = str(e) if str(e) else f"Error processing {detected_file_type.upper()} file"
        sys.stderr.write(f"❌ Error getting page count: {error_msg}\n")
        sys.stderr.write(f"   Exception type: {type(e).__name__}\n")
        import traceback
        traceback.print_exc(file=sys.stderr)
        sys.stderr.flush()
        raise HTTPException(status_code=400, detail=error_msg)
    
    return FileUploadResponse(
        upload_id=upload_id,
        page_count=page_count,
        file_size=file_size,
        temp_path=temp_path,
        file_type=detected_file_type,
        original_filename=file.filename or "unknown"
    )

# Keep old endpoint for backward compatibility
@router.post("/upload-pdf", response_model=FileUploadResponse)
async def upload_pdf_legacy(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user)
):
    """Legacy endpoint - redirects to upload-file"""
    return await upload_file(file, "question_paper", current_user)

@router.post("/submit-metadata")
async def submit_metadata(
    request: MetadataSubmitRequest,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Step 2: Submit metadata and start processing"""
    if current_user.role != "ADMIN":
        raise HTTPException(status_code=403, detail="Admin access required")
    
    # Check if temp file exists
    temp_path = os.path.join("tmp/uploads", request.upload_id)
    if not os.path.exists(temp_path):
        raise HTTPException(status_code=404, detail="Upload file not found")
    
    # Check for duplicate upload (same course_code + exam_type + exam_date)
    if request.file_type == "question_paper" and request.exam_type and request.exam_date:
        existing_paper = db.query(QuestionPaper).filter(
            QuestionPaper.course_code == request.course_code.upper(),
            QuestionPaper.exam_type == ExamType(request.exam_type),
            QuestionPaper.exam_date == request.exam_date
        ).first()
        
        if existing_paper:
            raise HTTPException(
                status_code=400,
                detail=f"A question paper for {request.course_code} with exam type {request.exam_type} on {request.exam_date} has already been uploaded (Paper ID: {existing_paper.paper_id})"
            )
    
    # Determine file extension from temp path
    file_ext = '.pdf'  # Default to PDF
    if os.path.exists(temp_path):
        _, ext = os.path.splitext(temp_path)
        if ext:
            file_ext = ext
    
    # Create permanent storage path
    permanent_dir = "storage/papers"
    os.makedirs(permanent_dir, exist_ok=True)
    permanent_path = os.path.join(permanent_dir, f"{request.upload_id}{file_ext}")
    
    # Move file to permanent storage
    shutil.move(temp_path, permanent_path)
    
    # Create database entry
    # For syllabus, exam_type might be None, so we need to handle that
    exam_type_enum = None
    if request.exam_type:
        exam_type_enum = ExamType(request.exam_type)
    elif request.file_type == "question_paper":
        raise HTTPException(status_code=400, detail="Exam type is required for question papers")
    
    paper = QuestionPaper(
        course_code=request.course_code,
        academic_year=request.academic_year,
        semester_type=SemesterType(request.semester_type),
        exam_type=exam_type_enum or ExamType.SEE,  # Default to SEE if not provided (for syllabus)
        exam_date=request.exam_date,
        pdf_path=permanent_path,
        temp_pdf_path=temp_path,
        uploaded_by=current_user.user_id,
        processing_status=ProcessingStatus.METADATA_PENDING,
        file_size=os.path.getsize(permanent_path)
    )
    
    db.add(paper)
    db.commit()
    db.refresh(paper)
    
    # For syllabus files, mark as completed immediately (no question extraction needed)
    # For question papers, process immediately (synchronously)
    if request.file_type == "syllabus":
        paper.processing_status = ProcessingStatus.COMPLETED
        paper.processing_progress = 100
        paper.total_questions_extracted = 0
        db.commit()
        task_id = None
    else:
        # Process question paper immediately (synchronously)
        try:
            # Import processing functions
            from app.tasks.processing import (
                convert_file_for_llm,
                extract_questions_with_llm,
                classify_questions_with_llm,
                detect_duplicates,
                save_questions
            )
            
            # Update status to processing
            paper.processing_status = ProcessingStatus.PROCESSING
            paper.processing_progress = 0
            db.commit()
            
            # Step 1: File Conversion (PDF/DOCX to text/images)
            paper.processing_progress = 10
            db.commit()
            file_content = convert_file_for_llm(paper)
            
            # Step 2: LLM Question Extraction
            paper.processing_progress = 30
            db.commit()
            questions = extract_questions_with_llm(file_content)
            
            # Step 3: LLM Classification (units and topic tags)
            paper.processing_progress = 50
            db.commit()
            classified_questions = classify_questions_with_llm(questions, paper.course_code, db)
            
            # Step 4: Duplicate Detection (simplified - paper-level duplicates already checked at upload)
            paper.processing_progress = 70
            db.commit()
            # Simple duplicate detection - just marks questions as canonical
            # Paper-level duplicate checking (same course, exam type, date) is done in submit_metadata
            deduplicated_questions = detect_duplicates(classified_questions, paper.course_code)
            
            # Step 5: Save to Database
            paper.processing_progress = 90
            db.commit()
            
            import logging
            logger = logging.getLogger(__name__)
            logger.info(f"About to save {len(deduplicated_questions)} questions for paper {paper.paper_id}")
            if deduplicated_questions:
                logger.info(f"Sample question data: {deduplicated_questions[0]}")
            
            save_questions(deduplicated_questions, paper, db)
            
            # Verify questions were saved
            saved_questions_count = db.query(Question).filter(Question.paper_id == paper.paper_id).count()
            logger.info(f"Total questions in database for paper {paper.paper_id}: {saved_questions_count}")
            
            # Update paper status to completed
            paper.processing_status = ProcessingStatus.COMPLETED
            paper.processing_progress = 100
            paper.total_questions_extracted = len(deduplicated_questions)
            db.commit()
            
            task_id = None  # No Celery task ID for synchronous processing
            
        except Exception as e:
            # Update paper status to failed
            paper.processing_status = ProcessingStatus.FAILED
            paper.processing_progress = 0
            db.commit()
            
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f"Failed to process question paper {paper.paper_id}: {str(e)}", exc_info=True)
            
            # Re-raise the exception so the API returns an error
            raise HTTPException(
                status_code=500,
                detail=f"Failed to process question paper: {str(e)}"
            )
    
    # Log activity
    from app.utils.activity_logger import log_activity
    file_type_label = "Question Paper" if request.file_type == "question_paper" else "Syllabus"
    log_activity(
        db=db,
        user_id=current_user.user_id,
        activity_type="PAPER_UPLOADED",
        description=f"Uploaded {file_type_label} for {request.course_code} ({request.academic_year}, {request.semester_type})",
        entity_type="question_paper",
        entity_id=str(paper.paper_id),
        metadata={"paper_id": paper.paper_id, "course_code": request.course_code, "file_type": request.file_type, "academic_year": request.academic_year, "semester_type": request.semester_type}
    )
    
    return {
        "paper_id": paper.paper_id,
        "task_id": task_id,
        "message": "Processing started" if request.file_type == "question_paper" else "Upload completed"
    }

@router.get("/processing-status/{paper_id}", response_model=ProcessingStatusResponse)
async def get_processing_status(
    paper_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Get processing status for a question paper"""
    if current_user.role != "ADMIN":
        raise HTTPException(status_code=403, detail="Admin access required")
    
    paper = db.query(QuestionPaper).filter(QuestionPaper.paper_id == paper_id).first()
    if not paper:
        raise HTTPException(status_code=404, detail="Question paper not found")
    
    return ProcessingStatusResponse(
        paper_id=paper.paper_id,
        status=paper.processing_status.value,
        progress=paper.processing_progress,
        questions_extracted=paper.total_questions_extracted
    )

@router.post("/retry-processing/{paper_id}")
async def retry_processing(
    paper_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Manually retry processing for a paper stuck in METADATA_PENDING"""
    if current_user.role != "ADMIN":
        raise HTTPException(status_code=403, detail="Admin access required")
    
    paper = db.query(QuestionPaper).filter(QuestionPaper.paper_id == paper_id).first()
    if not paper:
        raise HTTPException(status_code=404, detail="Question paper not found")
    
    if paper.processing_status != ProcessingStatus.METADATA_PENDING:
        raise HTTPException(
            status_code=400, 
            detail=f"Paper is not in METADATA_PENDING status. Current status: {paper.processing_status.value}"
        )
    
    try:
        # Reset status and try to queue the task again
        task = process_question_paper.delay(paper.paper_id)
        return {
            "paper_id": paper.paper_id,
            "task_id": task.id,
            "message": "Processing task queued successfully"
        }
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Failed to queue processing task. Make sure Celery worker is running: {str(e)}"
        )

@router.post("/questions/search")
async def search_questions(
    request: dict,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Search questions for admin question bank"""
    if current_user.role != "ADMIN":
        raise HTTPException(status_code=403, detail="Admin access required")
    
    from sqlalchemy.orm import joinedload
    
    # Build query
    query = db.query(Question).options(
        joinedload(Question.unit),
        joinedload(Question.paper),
        joinedload(Question.course)
    )
    
    # Apply filters
    filters = request.get("filters", {})
    
    if filters.get("course_codes"):
        query = query.filter(Question.course_code.in_(filters["course_codes"]))
    
    if filters.get("unit_ids"):
        query = query.filter(Question.unit_id.in_(filters["unit_ids"]))
    
    if filters.get("marks_min"):
        query = query.filter(Question.marks >= filters["marks_min"])
    
    if filters.get("marks_max"):
        query = query.filter(Question.marks <= filters["marks_max"])
    
    if filters.get("bloom_levels"):
        query = query.filter(Question.bloom_level.in_(filters["bloom_levels"]))
    
    if filters.get("exam_types"):
        query = query.join(QuestionPaper).filter(QuestionPaper.exam_type.in_(filters["exam_types"]))
    
    if filters.get("year_from"):
        query = query.join(QuestionPaper).filter(QuestionPaper.academic_year >= filters["year_from"])
    
    if filters.get("year_to"):
        query = query.join(QuestionPaper).filter(QuestionPaper.academic_year <= filters["year_to"])
    
    # Filter by review status
    if filters.get("review_status"):
        review_status = filters["review_status"]
        if review_status == "reviewed":
            query = query.filter(Question.is_reviewed == True)
        elif review_status == "non-reviewed":
            query = query.filter(Question.is_reviewed == False)
    
    # Text search
    search_query = request.get("query", "").strip()
    if search_query:
        query = query.filter(Question.question_text.ilike(f"%{search_query}%"))
    
    # Pagination
    page = request.get("page", 1)
    limit = request.get("limit", 20)
    offset = (page - 1) * limit
    
    total = query.count()
    questions = query.offset(offset).limit(limit).all()
    
    # Convert to response format
    import json
    results = []
    for q in questions:
        topic_tags = None
        if q.topic_tags:
            try:
                topic_tags = json.loads(q.topic_tags) if isinstance(q.topic_tags, str) else q.topic_tags
            except:
                topic_tags = None
        
        results.append({
            "question_id": q.question_id,
            "question_text": q.question_text,
            "marks": q.marks,
            "bloom_level": q.bloom_level.value if q.bloom_level else None,
            "bloom_category": q.bloom_category.value if q.bloom_category else None,
            "difficulty_level": q.difficulty_level.value if q.difficulty_level else None,
            "course_code": q.course_code,
            "unit_name": q.unit.unit_name if q.unit else None,
            "exam_type": q.paper.exam_type.value if q.paper else None,
            "academic_year": q.paper.academic_year if q.paper else None,
            "semester_type": q.paper.semester_type.value if q.paper else None,
            "exam_date": q.paper.exam_date.isoformat() if q.paper and q.paper.exam_date else None,
            "topic_tags": topic_tags,
            "is_reviewed": q.is_reviewed,
            "review_status": q.review_status.value if q.review_status else None,
            "image_path": q.image_path
        })
    
    return {
        "questions": results,
        "total": total,
        "page": page,
        "limit": limit
    }

@router.get("/review-queue", response_model=List[ReviewQueueItem])
async def get_review_queue(
    page: int = 1,
    limit: int = 20,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Get questions in review queue - includes all non-reviewed questions"""
    if current_user.role != "ADMIN":
        raise HTTPException(status_code=403, detail="Admin access required")
    
    from sqlalchemy.orm import joinedload
    
    offset = (page - 1) * limit
    
    # Get all non-reviewed questions
    non_reviewed_questions = db.query(Question).options(
        joinedload(Question.unit),
        joinedload(Question.paper),
        joinedload(Question.course)
    ).filter(
        Question.is_reviewed == False
    ).order_by(Question.created_at.desc()).offset(offset).limit(limit).all()
    
    # Get existing ReviewQueue entries for these questions
    question_ids = [q.question_id for q in non_reviewed_questions]
    existing_reviews = db.query(ReviewQueue).filter(
        ReviewQueue.question_id.in_(question_ids),
        ReviewQueue.status == "PENDING"
    ).all()
    
    # Create a map of question_id -> ReviewQueue entry
    review_map = {r.question_id: r for r in existing_reviews}
    
    # Build response items
    import json
    results = []
    for question in non_reviewed_questions:
        # Use existing ReviewQueue entry if available, otherwise create a virtual one
        if question.question_id in review_map:
            review_entry = review_map[question.question_id]
            suggested_correction = None
            if review_entry.suggested_correction:
                try:
                    suggested_correction = json.loads(review_entry.suggested_correction) if isinstance(review_entry.suggested_correction, str) else review_entry.suggested_correction
                except:
                    suggested_correction = {}
            
            # Parse unit topics
            unit_topics_list = None
            if question.unit and question.unit.topics:
                try:
                    unit_topics_list = json.loads(question.unit.topics) if isinstance(question.unit.topics, str) else question.unit.topics
                    if not isinstance(unit_topics_list, list):
                        # If not a list, try to split by comma
                        if isinstance(unit_topics_list, str):
                            unit_topics_list = [t.strip() for t in unit_topics_list.split(',') if t.strip()]
                        else:
                            unit_topics_list = []
                except:
                    unit_topics_list = None
            
            # Parse question topic tags
            topic_tags_list = None
            if question.topic_tags:
                try:
                    topic_tags_list = json.loads(question.topic_tags) if isinstance(question.topic_tags, str) else question.topic_tags
                except:
                    topic_tags_list = None
            
            results.append(ReviewQueueItem(
                review_id=review_entry.review_id,
                question_id=question.question_id,
                question_text=question.question_text,
                issue_type=review_entry.issue_type,
                suggested_correction=suggested_correction or {},
                priority=review_entry.priority,
                created_at=review_entry.created_at or question.created_at,
                marks=question.marks,
                bloom_level=question.bloom_level.value if question.bloom_level else None,
                bloom_category=question.bloom_category.value if question.bloom_category else None,
                unit_id=question.unit_id,
                unit_name=question.unit.unit_name if question.unit else None,
                unit_topics=unit_topics_list,
                course_code=question.course_code,
                course_name=question.course.course_name if question.course else None,
                topic_tags=topic_tags_list,
                image_path=question.image_path
            ))
        else:
            # Create a virtual ReviewQueueItem for questions without explicit review entries
            # Determine issue type based on question state
            issue_type = "NEEDS_REVIEW"
            if question.unit_id is None:
                issue_type = "AMBIGUOUS_UNIT"
            elif question.classification_confidence and question.classification_confidence < 0.7:
                issue_type = "LOW_CONFIDENCE"
            
            # Create a temporary review_id (negative to indicate virtual entry)
            # In production, you might want to actually create ReviewQueue entries for all non-reviewed questions
            
            # Parse unit topics
            unit_topics_list = None
            if question.unit and question.unit.topics:
                try:
                    unit_topics_list = json.loads(question.unit.topics) if isinstance(question.unit.topics, str) else question.unit.topics
                    if not isinstance(unit_topics_list, list):
                        # If not a list, try to split by comma
                        if isinstance(unit_topics_list, str):
                            unit_topics_list = [t.strip() for t in unit_topics_list.split(',') if t.strip()]
                        else:
                            unit_topics_list = []
                except:
                    unit_topics_list = None
            
            # Parse question topic tags
            topic_tags_list = None
            if question.topic_tags:
                try:
                    topic_tags_list = json.loads(question.topic_tags) if isinstance(question.topic_tags, str) else question.topic_tags
                except:
                    topic_tags_list = None
            
            results.append(ReviewQueueItem(
                review_id=-question.question_id,  # Virtual ID
                question_id=question.question_id,
                question_text=question.question_text,
                issue_type=issue_type,
                suggested_correction={
                    'unit_id': question.unit_id,
                    'unit_name': question.unit.unit_name if question.unit else None,
                    'bloom_level': question.bloom_level.value if question.bloom_level else None,
                    'bloom_category': question.bloom_category.value if question.bloom_category else None,
                    'marks': question.marks,
                    'topic_tags': topic_tags_list or []
                },
                priority=1 if question.unit_id is None else 2,
                created_at=question.created_at,
                marks=question.marks,
                bloom_level=question.bloom_level.value if question.bloom_level else None,
                bloom_category=question.bloom_category.value if question.bloom_category else None,
                unit_id=question.unit_id,
                unit_name=question.unit.unit_name if question.unit else None,
                unit_topics=unit_topics_list,
                course_code=question.course_code,
                course_name=question.course.course_name if question.course else None,
                topic_tags=topic_tags_list,
                image_path=question.image_path
            ))
    
    return results

@router.put("/approve-question/{question_id}")
async def approve_question(
    question_id: int,
    request: QuestionApprovalRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Approve or correct a question in review queue"""
    if current_user.role != "ADMIN":
        raise HTTPException(status_code=403, detail="Admin access required")
    
    question = db.query(Question).filter(Question.question_id == question_id).first()
    if not question:
        raise HTTPException(status_code=404, detail="Question not found")
    
    if request.approved:
        # Update question with corrections if provided
        if request.unit_id:
            question.unit_id = request.unit_id
        if request.bloom_level:
            question.bloom_level = BloomLevel(request.bloom_level)
        if request.marks:
            question.marks = request.marks
        
        # Mark question as reviewed
        question.is_reviewed = True
        question.review_status = ReviewStatus.APPROVED
        
        # Update or create review queue entry
        review_entry = db.query(ReviewQueue).filter(
            ReviewQueue.question_id == question_id
        ).first()
        
        if review_entry:
            review_entry.status = "APPROVED"
        else:
            # Create review queue entry if it doesn't exist (shouldn't happen with new code, but for backward compatibility)
            review_entry = ReviewQueue(
                question_id=question_id,
                issue_type="NEEDS_REVIEW",
                status="APPROVED",
                priority=3
            )
            db.add(review_entry)
    else:
        # Mark as rejected/corrected
        question.is_reviewed = True
        question.review_status = ReviewStatus.NEEDS_REVIEW
        
        # Update or create review queue entry
        review_entry = db.query(ReviewQueue).filter(
            ReviewQueue.question_id == question_id
        ).first()
        
        if review_entry:
            review_entry.status = "CORRECTED"
        else:
            # Create review queue entry if it doesn't exist
            review_entry = ReviewQueue(
                question_id=question_id,
                issue_type="NEEDS_REVIEW",
                status="CORRECTED",
                priority=3
            )
            db.add(review_entry)
    
    db.commit()
    
    # Log activity
    from app.utils.activity_logger import log_activity
    action = "approved" if request.approved else "rejected"
    log_activity(
        db=db,
        user_id=current_user.user_id,
        activity_type="REVIEW_RESOLVED",
        description=f"{action.capitalize()} question {question_id} from review queue",
        entity_type="question",
        entity_id=str(question_id),
        metadata={"question_id": question_id, "action": action, "approved": request.approved}
    )
    
    return {"message": "Question updated successfully"}

@router.get("/questions/{question_id}")
async def get_question(
    question_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Get question details by ID"""
    if current_user.role != "ADMIN":
        raise HTTPException(status_code=403, detail="Admin access required")
    
    from sqlalchemy.orm import joinedload
    
    question = db.query(Question).options(
        joinedload(Question.unit),
        joinedload(Question.course),
        joinedload(Question.paper)
    ).filter(Question.question_id == question_id).first()
    
    if not question:
        raise HTTPException(status_code=404, detail="Question not found")
    
    import json
    topic_tags = None
    if question.topic_tags:
        try:
            topic_tags = json.loads(question.topic_tags) if isinstance(question.topic_tags, str) else question.topic_tags
        except:
            topic_tags = None
    
    return {
        "question_id": question.question_id,
        "question_text": question.question_text,
        "question_number": question.question_number,
        "marks": question.marks,
        "bloom_level": question.bloom_level.value if question.bloom_level else None,
        "bloom_category": question.bloom_category.value if question.bloom_category else None,
        "unit_id": question.unit_id,
        "unit_name": question.unit.unit_name if question.unit else None,
        "course_code": question.course_code,
        "topic_tags": topic_tags,
        "is_reviewed": question.is_reviewed,
        "review_status": question.review_status.value if question.review_status else None,
        "image_path": question.image_path,
        "paper_id": question.paper_id
    }

@router.put("/questions/{question_id}")
async def update_question(
    question_id: int,
    request: QuestionUpdateRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Update question metadata (unit, topic tags, marks, Bloom's taxonomy)"""
    if current_user.role != "ADMIN":
        raise HTTPException(status_code=403, detail="Admin access required")
    
    question = db.query(Question).filter(Question.question_id == question_id).first()
    if not question:
        raise HTTPException(status_code=404, detail="Question not found")
    
    # Update fields if provided
    if request.unit_id is not None:
        question.unit_id = request.unit_id
    if request.marks is not None:
        question.marks = request.marks
    if request.bloom_taxonomy_level is not None:
        question.bloom_level = BloomLevel(request.bloom_taxonomy_level)
    if request.bloom_category is not None:
        question.bloom_category = BloomCategory(request.bloom_category)
    if request.question_text is not None:
        question.question_text = request.question_text
    if request.topic_tags is not None:
        import json
        question.topic_tags = json.dumps(request.topic_tags)
    
    # Don't automatically mark as reviewed when editing
    # Admin should explicitly approve questions to mark them as reviewed
    # This allows editing questions in the review queue without removing them
    
    db.commit()
    db.refresh(question)
    
    # Log activity
    from app.utils.activity_logger import log_activity
    log_activity(
        db=db,
        user_id=current_user.user_id,
        activity_type="QUESTION_UPDATED",
        description=f"Updated question {question_id}",
        entity_type="question",
        entity_id=str(question_id),
        metadata={"question_id": question_id, "updates": request.dict(exclude_unset=True)}
    )
    
    return {"message": "Question updated successfully", "question_id": question_id}

@router.post("/questions/{question_id}/image")
async def upload_question_image(
    question_id: int,
    file: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Upload or replace image for a question (diagram, figure, etc.)"""
    if current_user.role != "ADMIN":
        raise HTTPException(status_code=403, detail="Admin access required")
    
    question = db.query(Question).filter(Question.question_id == question_id).first()
    if not question:
        raise HTTPException(status_code=404, detail="Question not found")
    
    # Validate file type
    if not file.content_type or not file.content_type.startswith('image/'):
        raise HTTPException(status_code=400, detail="File must be an image")
    
    # Create directory for question images
    question_images_dir = os.path.join("storage", "question_images", str(question_id))
    os.makedirs(question_images_dir, exist_ok=True)
    
    # Save image
    file_ext = os.path.splitext(file.filename)[1] or '.png'
    image_filename = f"question_image{file_ext}"
    image_path = os.path.join(question_images_dir, image_filename)
    
    with open(image_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    
    # Update question with image path
    question.image_path = image_path
    db.commit()
    
    # Log activity
    from app.utils.activity_logger import log_activity
    log_activity(
        db=db,
        user_id=current_user.user_id,
        activity_type="QUESTION_IMAGE_UPLOADED",
        description=f"Uploaded image for question {question_id}",
        entity_type="question",
        entity_id=str(question_id),
        metadata={"question_id": question_id, "image_path": image_path}
    )
    
    return {"message": "Image uploaded successfully", "image_path": image_path}

@router.post("/questions/{question_id}/approve")
async def approve_question_direct(
    question_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Approve a question (mark as reviewed)"""
    if current_user.role != "ADMIN":
        raise HTTPException(status_code=403, detail="Admin access required")
    
    question = db.query(Question).filter(Question.question_id == question_id).first()
    if not question:
        raise HTTPException(status_code=404, detail="Question not found")
    
    # Mark as reviewed and approved
    question.is_reviewed = True
    question.review_status = ReviewStatus.APPROVED
    
    # Update review queue if exists
    review_entry = db.query(ReviewQueue).filter(ReviewQueue.question_id == question_id).first()
    if review_entry:
        review_entry.status = "APPROVED"
    
    db.commit()
    
    # Log activity
    from app.utils.activity_logger import log_activity
    log_activity(
        db=db,
        user_id=current_user.user_id,
        activity_type="QUESTION_APPROVED",
        description=f"Approved question {question_id}",
        entity_type="question",
        entity_id=str(question_id),
        metadata={"question_id": question_id}
    )
    
    return {"message": "Question approved successfully"}

@router.get("/analytics/dashboard")
async def get_analytics_dashboard(
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Get analytics dashboard data"""
    try:
        if current_user.role != "ADMIN":
            raise HTTPException(status_code=403, detail="Admin access required")
        
        # Get basic statistics
        total_papers = db.query(QuestionPaper).count()
        total_questions = db.query(Question).count()
        pending_reviews = db.query(ReviewQueue).filter(ReviewQueue.status == "PENDING").count()
        
        # Get Bloom taxonomy distribution (handle empty database)
        try:
            bloom_results = db.query(Question.bloom_category, func.count(Question.question_id)).group_by(Question.bloom_category).all()
            bloom_distribution = {str(category or "Unknown"): count for category, count in bloom_results}
        except Exception as e:
            sys.stderr.write(f"⚠️  Error getting bloom distribution: {e}\n")
            bloom_distribution = {}
        
        # Get course-wise breakdown (handle empty database)
        try:
            course_results = db.query(QuestionPaper.course_code, func.count(QuestionPaper.paper_id)).group_by(QuestionPaper.course_code).all()
            course_breakdown = {str(code or "Unknown"): count for code, count in course_results}
        except Exception as e:
            sys.stderr.write(f"⚠️  Error getting course breakdown: {e}\n")
            course_breakdown = {}
        
        return {
            "total_papers": total_papers,
            "total_questions": total_questions,
            "pending_reviews": pending_reviews,
            "bloom_distribution": bloom_distribution,
            "course_breakdown": course_breakdown
        }
    except Exception as e:
        import traceback
        sys.stderr.write(f"❌ Error in get_analytics_dashboard: {e}\n")
        traceback.print_exc(file=sys.stderr)
        sys.stderr.flush()
        raise HTTPException(status_code=500, detail=f"Error fetching dashboard data: {str(e)}")

class ActivityLogResponse(BaseModel):
    log_id: int
    user_id: int
    username: str
    activity_type: str
    entity_type: Optional[str] = None
    entity_id: Optional[str] = None
    description: str
    metadata: Optional[dict] = None  # Frontend expects 'metadata', we'll map from activity_metadata
    created_at: datetime
    
    class Config:
        from_attributes = True

@router.get("/activity-logs")
async def get_activity_logs(
    page: int = Query(1, ge=1),
    limit: int = Query(50, ge=1, le=100),
    activity_type: Optional[str] = Query(None),
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Get activity logs (admin only)"""
    if current_user.role != "ADMIN":
        raise HTTPException(status_code=403, detail="Admin access required")
    
    from app.models.activity_log import ActivityLog
    from sqlalchemy import desc
    from sqlalchemy.orm import joinedload
    
    query = db.query(ActivityLog).options(joinedload(ActivityLog.user))
    
    if activity_type:
        query = query.filter(ActivityLog.activity_type == activity_type)
    
    offset = (page - 1) * limit
    logs = query.order_by(desc(ActivityLog.created_at)).offset(offset).limit(limit).all()
    total = query.count()
    
    # Format response with username
    log_responses = []
    for log in logs:
        log_responses.append(ActivityLogResponse(
            log_id=log.log_id,
            user_id=log.user_id,
            username=log.user.username if log.user else "Unknown",
            activity_type=log.activity_type,
            entity_type=log.entity_type,
            entity_id=log.entity_id,
            description=log.description,
            metadata=log.activity_metadata,  # Map activity_metadata to metadata for frontend
            created_at=log.created_at
        ))
    
    return {
        "logs": log_responses,
        "total": total,
        "page": page,
        "limit": limit,
        "pages": (total + limit - 1) // limit
    }

class PaperResponse(BaseModel):
    paper_id: int
    course_code: str
    course_name: Optional[str] = None
    academic_year: int
    semester_type: str
    exam_type: str
    exam_date: Optional[datetime] = None
    pdf_path: Optional[str] = None
    uploaded_by: int
    uploader_username: Optional[str] = None
    processing_status: str
    processing_progress: float
    total_questions_extracted: int
    file_size: Optional[int] = None
    page_count: Optional[int] = None
    created_at: datetime
    
    class Config:
        from_attributes = True

class PaperUpdate(BaseModel):
    course_code: Optional[str] = None
    academic_year: Optional[int] = None
    semester_type: Optional[str] = None
    exam_type: Optional[str] = None
    exam_date: Optional[datetime] = None

@router.get("/uploads", response_model=List[PaperResponse])
async def get_all_uploads(
    page: int = Query(1, ge=1),
    limit: int = Query(50, ge=1, le=100),
    course_code: Optional[str] = Query(None),
    file_type: Optional[str] = Query(None),  # 'question_paper' or 'syllabus' - determined by exam_type
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Get all uploads (question papers and syllabi) - admin only"""
    if current_user.role != "ADMIN":
        raise HTTPException(status_code=403, detail="Admin access required")
    
    from sqlalchemy.orm import joinedload
    from sqlalchemy import desc
    
    query = db.query(QuestionPaper).options(joinedload(QuestionPaper.uploader), joinedload(QuestionPaper.course))
    
    # Filter by course if provided
    if course_code:
        query = query.filter(QuestionPaper.course_code == course_code.upper())
    
    # Filter by file type (syllabus = SEE exam type, question paper = other exam types)
    if file_type == "syllabus":
        query = query.filter(QuestionPaper.exam_type == ExamType.SEE)
    elif file_type == "question_paper":
        query = query.filter(QuestionPaper.exam_type != ExamType.SEE)
    
    offset = (page - 1) * limit
    papers = query.order_by(desc(QuestionPaper.created_at)).offset(offset).limit(limit).all()
    total = query.count()
    
    # Format response
    paper_responses = []
    for paper in papers:
        paper_responses.append(PaperResponse(
            paper_id=paper.paper_id,
            course_code=paper.course_code,
            course_name=paper.course.course_name if paper.course else None,
            academic_year=paper.academic_year,
            semester_type=paper.semester_type.value,
            exam_type=paper.exam_type.value,
            exam_date=paper.exam_date,
            pdf_path=paper.pdf_path,
            uploaded_by=paper.uploaded_by,
            uploader_username=paper.uploader.username if paper.uploader else None,
            processing_status=paper.processing_status.value,
            processing_progress=paper.processing_progress,
            total_questions_extracted=paper.total_questions_extracted,
            file_size=paper.file_size,
            page_count=paper.page_count,
            created_at=paper.created_at
        ))
    
    return paper_responses

@router.get("/uploads/{paper_id}", response_model=PaperResponse)
async def get_upload(
    paper_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Get specific upload details - admin only"""
    if current_user.role != "ADMIN":
        raise HTTPException(status_code=403, detail="Admin access required")
    
    from sqlalchemy.orm import joinedload
    
    paper = db.query(QuestionPaper).options(
        joinedload(QuestionPaper.uploader),
        joinedload(QuestionPaper.course)
    ).filter(QuestionPaper.paper_id == paper_id).first()
    
    if not paper:
        raise HTTPException(status_code=404, detail="Upload not found")
    
    return PaperResponse(
        paper_id=paper.paper_id,
        course_code=paper.course_code,
        course_name=paper.course.course_name if paper.course else None,
        academic_year=paper.academic_year,
        semester_type=paper.semester_type.value,
        exam_type=paper.exam_type.value,
        exam_date=paper.exam_date,
        pdf_path=paper.pdf_path,
        uploaded_by=paper.uploaded_by,
        uploader_username=paper.uploader.username if paper.uploader else None,
        processing_status=paper.processing_status.value,
        processing_progress=paper.processing_progress,
        total_questions_extracted=paper.total_questions_extracted,
        file_size=paper.file_size,
        page_count=paper.page_count,
        created_at=paper.created_at
    )

@router.get("/uploads/{paper_id}/questions")
async def get_paper_questions(
    paper_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Get all questions from a specific paper - admin only"""
    if current_user.role != "ADMIN":
        raise HTTPException(status_code=403, detail="Admin access required")
    
    # Verify paper exists
    paper = db.query(QuestionPaper).filter(QuestionPaper.paper_id == paper_id).first()
    if not paper:
        raise HTTPException(status_code=404, detail="Question paper not found")
    
    # Get all questions for this paper
    questions = db.query(Question).filter(Question.paper_id == paper_id).order_by(Question.question_number).all()
    
    import json
    results = []
    for q in questions:
        # Parse topic tags
        topic_tags = None
        if q.topic_tags:
            try:
                topic_tags = json.loads(q.topic_tags) if isinstance(q.topic_tags, str) else q.topic_tags
            except:
                topic_tags = None
        
        results.append({
            "question_id": q.question_id,
            "question_number": q.question_number,
            "question_text": q.question_text,
            "marks": q.marks,
            "bloom_level": q.bloom_level.value if q.bloom_level else None,
            "bloom_category": q.bloom_category.value if q.bloom_category else None,
            "unit_id": q.unit_id,
            "unit_name": q.unit.unit_name if q.unit else None,
            "topic_tags": topic_tags,
            "classification_confidence": q.classification_confidence,
            "is_reviewed": q.is_reviewed,
            "review_status": q.review_status.value if q.review_status else None,
            "is_canonical": q.is_canonical,
            "parent_question_id": q.parent_question_id,
            "page_number": q.page_number,
            "created_at": q.created_at.isoformat() if q.created_at else None
        })
    
    return {
        "paper_id": paper_id,
        "course_code": paper.course_code,
        "total_questions": len(results),
        "questions": results
    }

@router.put("/uploads/{paper_id}", response_model=PaperResponse)
async def update_upload(
    paper_id: int,
    paper_update: PaperUpdate,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Update upload metadata - admin only"""
    if current_user.role != "ADMIN":
        raise HTTPException(status_code=403, detail="Admin access required")
    
    paper = db.query(QuestionPaper).filter(QuestionPaper.paper_id == paper_id).first()
    if not paper:
        raise HTTPException(status_code=404, detail="Upload not found")
    
    # Update fields if provided
    if paper_update.course_code is not None:
        # Verify course exists
        course = db.query(Course).filter(Course.course_code == paper_update.course_code.upper()).first()
        if not course:
            raise HTTPException(status_code=404, detail=f"Course {paper_update.course_code} not found")
        paper.course_code = paper_update.course_code.upper()
    
    if paper_update.academic_year is not None:
        paper.academic_year = paper_update.academic_year
    
    if paper_update.semester_type is not None:
        if paper_update.semester_type.upper() not in ['ODD', 'EVEN']:
            raise HTTPException(status_code=400, detail="Semester type must be ODD or EVEN")
        paper.semester_type = SemesterType(paper_update.semester_type.upper())
    
    if paper_update.exam_type is not None:
        try:
            paper.exam_type = ExamType(paper_update.exam_type)
        except ValueError:
            raise HTTPException(status_code=400, detail=f"Invalid exam type: {paper_update.exam_type}")
    
    if paper_update.exam_date is not None:
        paper.exam_date = paper_update.exam_date
    
    db.commit()
    db.refresh(paper)
    
    # Log activity
    from app.utils.activity_logger import log_activity
    log_activity(
        db=db,
        user_id=current_user.user_id,
        activity_type="PAPER_UPDATED",
        description=f"Updated upload {paper_id} for {paper.course_code}",
        entity_type="question_paper",
        entity_id=str(paper_id),
        metadata={"paper_id": paper_id, "course_code": paper.course_code}
    )
    
    # Reload with relationships
    from sqlalchemy.orm import joinedload
    db.refresh(paper)
    paper = db.query(QuestionPaper).options(
        joinedload(QuestionPaper.uploader),
        joinedload(QuestionPaper.course)
    ).filter(QuestionPaper.paper_id == paper_id).first()
    
    return PaperResponse(
        paper_id=paper.paper_id,
        course_code=paper.course_code,
        course_name=paper.course.course_name if paper.course else None,
        academic_year=paper.academic_year,
        semester_type=paper.semester_type.value,
        exam_type=paper.exam_type.value,
        exam_date=paper.exam_date,
        pdf_path=paper.pdf_path,
        uploaded_by=paper.uploaded_by,
        uploader_username=paper.uploader.username if paper.uploader else None,
        processing_status=paper.processing_status.value,
        processing_progress=paper.processing_progress,
        total_questions_extracted=paper.total_questions_extracted,
        file_size=paper.file_size,
        page_count=paper.page_count,
        created_at=paper.created_at
    )

@router.delete("/uploads/{paper_id}")
async def delete_upload(
    paper_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    """Delete an upload and all associated questions - admin only"""
    if current_user.role != "ADMIN":
        raise HTTPException(status_code=403, detail="Admin access required")
    
    paper = db.query(QuestionPaper).filter(QuestionPaper.paper_id == paper_id).first()
    if not paper:
        raise HTTPException(status_code=404, detail="Upload not found")
    
    course_code = paper.course_code
    pdf_path = paper.pdf_path
    
    # Delete associated questions
    from app.models.question import Question
    db.query(Question).filter(Question.paper_id == paper_id).delete()
    
    # Delete the paper
    db.delete(paper)
    db.commit()
    
    # Delete the file from storage
    if pdf_path and os.path.exists(pdf_path):
        try:
            os.remove(pdf_path)
        except Exception as e:
            sys.stderr.write(f"Warning: Failed to delete file {pdf_path}: {e}\n")
    
    # Log activity
    from app.utils.activity_logger import log_activity
    log_activity(
        db=db,
        user_id=current_user.user_id,
        activity_type="PAPER_DELETED",
        description=f"Deleted upload {paper_id} for {course_code}",
        entity_type="question_paper",
        entity_id=str(paper_id),
        metadata={"paper_id": paper_id, "course_code": course_code}
    )
    
    return {"message": f"Upload {paper_id} deleted successfully"}
