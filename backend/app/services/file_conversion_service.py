"""
File Conversion Service
Converts PDF and DOCX files to text and images for LLM processing
"""
import os
import base64
from typing import Dict, List, Optional
from PIL import Image
import io

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    try:
        import PyPDF2
        PYPDF2_AVAILABLE = True
    except ImportError:
        PYPDF2_AVAILABLE = False

try:
    from docx import Document
    from docx.document import Document as DocumentType
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from pdf2image import convert_from_path
    PDF2IMAGE_AVAILABLE = True
except ImportError:
    PDF2IMAGE_AVAILABLE = False


class FileConversionService:
    """Service to convert PDF/DOCX files to text and images for LLM processing"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc']
    
    def convert_file(self, file_path: str) -> Dict:
        """
        Convert file (PDF or DOCX) to text and images
        
        Returns:
            Dict with:
            - text: Extracted text content
            - images: List of base64-encoded images
            - file_type: 'pdf' or 'docx'
            - page_count: Number of pages
        """
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            return self.convert_pdf_to_text_and_images(file_path)
        elif file_ext in ['.docx', '.doc']:
            return self.convert_docx_to_text_and_images(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    def convert_pdf_to_text_and_images(self, pdf_path: str) -> Dict:
        """Convert PDF to text and images"""
        result = {
            'text': '',
            'images': [],
            'file_type': 'pdf',
            'page_count': 0,
            'pages': []
        }
        
        # Try to extract text first (for text-based PDFs)
        text_content = self._extract_text_from_pdf(pdf_path)
        
        # Convert PDF pages to images (for vision models and image-based PDFs)
        page_images = self._convert_pdf_pages_to_images(pdf_path)
        
        result['page_count'] = len(page_images)
        
        # Combine text from all pages
        if text_content:
            result['text'] = text_content
        else:
            # If no text extracted, we'll rely on images for OCR/LLM vision
            result['text'] = ""
        
        # Convert images to base64 for LLM API
        for i, img in enumerate(page_images):
            img_base64 = self._image_to_base64(img)
            result['images'].append({
                'page_number': i + 1,
                'base64': img_base64,
                'format': 'png'
            })
            result['pages'].append({
                'page_number': i + 1,
                'text': '',  # Will be extracted by LLM from image
                'image': img_base64
            })
        
        return result
    
    def _extract_text_from_pdf(self, pdf_path: str) -> str:
        """Extract text from PDF using pdfplumber or PyPDF2"""
        text_parts = []
        
        if PDFPLUMBER_AVAILABLE:
            try:
                with pdfplumber.open(pdf_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
            except Exception as e:
                print(f"Error extracting text with pdfplumber: {e}")
        
        elif PYPDF2_AVAILABLE:
            try:
                with open(pdf_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
            except Exception as e:
                print(f"Error extracting text with PyPDF2: {e}")
        
        return '\n\n'.join(text_parts)
    
    def _convert_pdf_pages_to_images(self, pdf_path: str) -> List[Image.Image]:
        """Convert PDF pages to PIL Images"""
        if not PDF2IMAGE_AVAILABLE:
            raise ImportError("pdf2image is required for PDF to image conversion. Install with: pip install pdf2image")
        
        try:
            # Convert PDF to images (300 DPI for good quality)
            images = convert_from_path(pdf_path, dpi=300)
            return images
        except Exception as e:
            raise Exception(f"Failed to convert PDF to images: {e}")
    
    def convert_docx_to_text_and_images(self, docx_path: str) -> Dict:
        """Convert DOCX to text and images"""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX processing. Install with: pip install python-docx")
        
        result = {
            'text': '',
            'images': [],
            'file_type': 'docx',
            'page_count': 0,
            'pages': []
        }
        
        try:
            doc = Document(docx_path)
            
            # Extract text
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join([cell.text for cell in row.cells])
                    if row_text.strip():
                        text_parts.append(row_text)
            
            result['text'] = '\n\n'.join(text_parts)
            
            # Extract embedded images
            image_count = 0
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        image_part = rel.target_part
                        image_bytes = image_part.blob
                        img = Image.open(io.BytesIO(image_bytes))
                        img_base64 = self._image_to_base64(img)
                        image_count += 1
                        result['images'].append({
                            'page_number': image_count,
                            'base64': img_base64,
                            'format': img.format.lower() if img.format else 'png'
                        })
                    except Exception as e:
                        print(f"Error extracting image from DOCX: {e}")
                        continue
            
            result['page_count'] = 1  # DOCX doesn't have pages, but we count images
            result['pages'].append({
                'page_number': 1,
                'text': result['text'],
                'images': result['images']
            })
            
        except Exception as e:
            raise Exception(f"Failed to process DOCX file: {e}")
        
        return result
    
    def _image_to_base64(self, image: Image.Image) -> str:
        """Convert PIL Image to base64 string"""
        buffered = io.BytesIO()
        # Convert to RGB if necessary (for PNG/JPG compatibility)
        if image.mode != 'RGB':
            image = image.convert('RGB')
        image.save(buffered, format='PNG')
        img_str = base64.b64encode(buffered.getvalue()).decode('utf-8')
        return img_str
    
    def prepare_for_llm(self, conversion_result: Dict, max_images: int = 10) -> Dict:
        """
        Prepare content for LLM API
        
        Args:
            conversion_result: Result from convert_file()
            max_images: Maximum number of images to include (to avoid token limits)
        
        Returns:
            Dict formatted for OpenAI API
        """
        # Limit number of images
        images = conversion_result['images'][:max_images]
        
        # Prepare content array for OpenAI API
        content = []
        
        # Add text if available
        if conversion_result.get('text'):
            content.append({
                "type": "text",
                "text": conversion_result['text']
            })
        
        # Add images
        for img_data in images:
            content.append({
                "type": "image_url",
                "image_url": {
                    "url": f"data:image/{img_data['format']};base64,{img_data['base64']}"
                }
            })
        
        return {
            "text": conversion_result.get('text', ''),
            "content": content,
            "page_count": conversion_result.get('page_count', 0),
            "has_images": len(images) > 0
        }

